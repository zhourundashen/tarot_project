from docx import Document
import json

docx_path = "C:/Users/MAKOTO/Desktop/tarot_project/tarot_books/5.塔罗葵花宝典（精排版）.docx"

print("Reading Word document...")
doc = Document(docx_path)

content = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        content.append(text)

full_text = "\n".join(content)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Non-empty paragraphs: {len(content)}")
print(f"\nFirst 3000 chars:\n{'-'*50}")
print(full_text[:3000])

output_path = "C:/Users/MAKOTO/Desktop/tarot_project/knowledge/books/raw/sunflower_treasure_docx.txt"
with open(output_path, 'w', encoding='utf-8') as f:
    f.write(full_text)
print(f"\n\nSaved to: {output_path}")
