import os
from docx import Document

def get_tarot_data():
    """返回塔罗牌数据"""
    return {
        "major": [
            {
                "num": "0号", 
                "name": "愚人（The Fool）",
                "image": "韦特牌的愚人画面里，一个年轻人站在悬崖边，背着包袱，手持白玫瑰，准备踏上旅程。他的眼神充满好奇与期待，脚下是万丈深渊，但他毫不畏惧，仿佛在告诉我们：人生就是一场冒险，要勇敢地迈出第一步。",
                "upright": "新的开始、无限可能、纯真无畏、活在当下、自由自在。预示着一个新的开始即将到来，你将踏上全新的旅程，充满无限的可能性。这是一个充满希望和机遇的时刻，鼓励你保持初心，勇敢前行。",
                "reverse_neg": "鲁莽冲动、缺乏计划、逃避责任、不切实际、危险冒险。你可能会过于冲动，缺乏周密的计划就贸然行动，或者逃避应有的责任，做出不切实际的决定。",
                "reverse_pos": "在保持初心的同时，加入理性的思考，制定切实可行的计划，承担责任，让冒险变得更加安全和有意义。"
            },
            {
                "num": "1号", 
                "name": "魔术师（The Magician）",
                "image": "韦特牌的魔术师画面里，一位身着红色衣服的男子站在石桌前，桌上摆放着代表四大元素的法器：权杖、圣杯、宝剑、星币。他一手指向天空，一手指向大地，象征着天地之间的连接。",
                "upright": "创造力与执行力、资源整合、沟通表达、心想事成、技能展现。你拥有将想法转化为现实的强大能力，能够巧妙地整合各种资源，通过有效的沟通实现目标。",
                "reverse_neg": "创造力枯竭、执行力缺失、沟通不畅、资源失控、投机取巧、谎言与欺骗、能力不足、想法与行动脱节。",
                "reverse_pos": "放下不切实际的空想，脚踏实地打磨自己的能力，将想法转化为具体的行动；梳理自己的思路，清晰地表达自己的意图，改善沟通方式，获得他人的支持。"
            }
        ]
    }

def create_tarot_document():
    """创建塔罗牌解读文档"""
    print("开始创建塔罗牌文档...")
    
    # 创建文档对象
    doc = Document()
    print("文档对象创建成功")
    
    # 添加介绍段落
    intro_text = (
        '本解读以韦特塔罗牌经典牌面为基础，为每一张牌提供牌面细节解读、'
        '核心原型定位、心理层面意义、正位核心解读、'
        '逆位核心解读（含消极走向与积极走向），既保留塔罗牌的经典象征意义，'
        '又结合现代心理视角，让解读更贴合当代人的生命体验。'
        '塔罗牌的核心并非预测未来，而是照见内心，通过牌面的象征唤醒自我认知，'
        '引导我们做出符合内心的选择，最终成为完整、真实、喜悦的自己。'
    )
    doc.add_paragraph(intro_text, style='Normal')
    print("介绍段落添加完成")

    data = get_tarot_data()
    print(f"获取到 {len(data['major'])} 张塔罗牌数据")

    # 添加大阿卡纳部分标题
    doc.add_paragraph('第一部分：22张大阿卡纳——愚人之旅的灵魂成长轨迹', style='Heading 1')
    
    # 遍历大阿卡纳牌
    for i, card in enumerate(data["major"]):
        print(f"处理第 {i+1} 张牌: {card['name']}")
        # 添加牌名标题
        doc.add_paragraph(f"{card['num']} {card['name']}", style='Heading 2')
        
        # 添加牌面解读
        doc.add_paragraph('【牌面解读】', style='Heading 3')
        doc.add_paragraph(card["image"], style='Normal')
        
        # 添加正位解读
        doc.add_paragraph('【正位核心解读】', style='Heading 3')
        doc.add_paragraph(card["upright"], style='Normal')
        
        # 添加逆位解读
        doc.add_paragraph('【逆位核心解读】', style='Heading 3')
        doc.add_paragraph('消极走向：' + card["reverse_neg"], style='Normal')
        doc.add_paragraph('积极走向：' + card["reverse_pos"], style='Normal')

    # 确保在当前目录保存文件
    file_path = os.path.join(os.getcwd(), 'tarot_interpretation.docx')
    print(f"准备保存文件到: {file_path}")
    
    # 保存文档
    doc.save(file_path)
    print("塔罗牌解读文档已生成！")
    
    # 验证文件是否存在
    if os.path.exists(file_path):
        file_size = os.path.getsize(file_path)
        print(f"文件生成成功！路径: {file_path}, 大小: {file_size} 字节")
    else:
        print("文件生成失败！")

# 如果直接运行此脚本
if __name__ == "__main__":
    print(f"当前工作目录: {os.getcwd()}")
    create_tarot_document()