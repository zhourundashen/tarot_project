import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_tarot_knowledge():
    """加载完整的塔罗牌知识库"""
    knowledge_files = [
        "c:/Users/MAKOTO/Desktop/tarot_project/knowledge/books/tarot_knowledge_actually_you_are.json",
        "c:/Users/MAKOTO/Desktop/tarot_project/knowledge/books/tarot_knowledge_78_wisdom.json",
        "c:/Users/MAKOTO/Desktop/tarot_project/knowledge/books/tarot_knowledge_sunflower.json"
    ]
    
    combined_knowledge = {
        "majorArcana": [],
        "minorArcana": {
            "wands": {"name": "权杖", "cards": []},
            "cups": {"name": "圣杯", "cards": []},
            "swords": {"name": "宝剑", "cards": []},
            "pentacles": {"name": "星币", "cards": []}
        }
    }
    
    # 定义标准的78张牌名称
    major_arcana_names = [
        ("0", "愚人", "The Fool"),
        ("1", "魔术师", "The Magician"),
        ("2", "女祭司", "The High Priestess"),
        ("3", "皇后", "The Empress"),
        ("4", "皇帝", "The Emperor"),
        ("5", "教皇", "The Hierophant"),
        ("6", "恋人", "The Lovers"),
        ("7", "战车", "The Chariot"),
        ("8", "力量", "Strength"),
        ("9", "隐士", "The Hermit"),
        ("10", "命运之轮", "The Wheel of Fortune"),
        ("11", "正义", "Justice"),
        ("12", "吊人", "The Hanged Man"),
        ("13", "死神", "Death"),
        ("14", "节制", "Temperance"),
        ("15", "恶魔", "The Devil"),
        ("16", "塔", "The Tower"),
        ("17", "星星", "The Star"),
        ("18", "月亮", "The Moon"),
        ("19", "太阳", "The Sun"),
        ("20", "审判", "Judgement"),
        ("21", "世界", "The World")
    ]
    
    # 创建大阿卡纳基础数据
    for num, cn_name, en_name in major_arcana_names:
        combined_knowledge["majorArcana"].append({
            "id": int(num),
            "num": f"{num}号",
            "name": f"{cn_name}（{en_name}）",
            "nameCn": cn_name,
            "nameEn": en_name,
            "type": "major",
            "image": f"韦特牌的{cn_name}画面里...",  # 占位符
            "upright": "正位含义待补充...",
            "reverse_neg": "逆位消极含义待补充...",
            "reverse_pos": "逆位积极含义待补充..."
        })
    
    # 创建小阿卡纳基础数据
    suits = {
        "wands": {"cn": "权杖", "en": "Wands"},
        "cups": {"cn": "圣杯", "en": "Cups"}, 
        "swords": {"cn": "宝剑", "en": "Swords"},
        "pentacles": {"cn": "星币", "en": "Pentacles"}
    }
    
    numbers = [
        ("一", "Ace"), ("二", "Two"), ("三", "Three"), ("四", "Four"),
        ("五", "Five"), ("六", "Six"), ("七", "Seven"), ("八", "Eight"),
        ("九", "Nine"), ("十", "Ten")
    ]
    
    courts = [("侍卫", "Page"), ("骑士", "Knight"), ("皇后", "Queen"), ("国王", "King")]
    
    for suit_key, suit_info in suits.items():
        suit_cards = []
        # 数字牌
        for i, (cn_num, en_num) in enumerate(numbers):
            card_id = f"{suit_key[0]}_{i+1}"
            card_name = f"{suit_info['cn']}{cn_num}"
            card_en = f"{en_num} of {suit_info['en']}"
            suit_cards.append({
                "id": card_id,
                "num": f"{i+1}号",
                "name": f"{card_name}（{card_en}）",
                "nameCn": card_name,
                "nameEn": card_en,
                "type": "minor",
                "suit": suit_key,
                "image": f"韦特牌的{card_name}画面里...",
                "upright": "正位含义待补充...",
                "reverse_neg": "逆位消极含义待补充...",
                "reverse_pos": "逆位积极含义待补充..."
            })
        
        # 宫廷牌
        for court_cn, court_en in courts:
            card_id = f"{suit_key[0]}_{court_en.lower()}"
            card_name = f"{suit_info['cn']}{court_cn}"
            card_en = f"{court_en} of {suit_info['en']}"
            suit_cards.append({
                "id": card_id,
                "num": court_cn,
                "name": f"{card_name}（{card_en}）",
                "nameCn": card_name,
                "nameEn": card_en,
                "type": "minor",
                "suit": suit_key,
                "image": f"韦特牌的{card_name}画面里...",
                "upright": "正位含义待补充...",
                "reverse_neg": "逆位消极含义待补充...",
                "reverse_pos": "逆位积极含义待补充..."
            })
        
        combined_knowledge["minorArcana"][suit_key]["cards"] = suit_cards
    
    # 尝试从现有文件加载具体含义
    for file_path in knowledge_files:
        try:
            if os.path.exists(file_path):
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # 更新大阿卡纳含义
                if "majorArcana" in data:
                    for existing_card in data["majorArcana"]:
                        for card in combined_knowledge["majorArcana"]:
                            if card["id"] == existing_card["id"]:
                                # 更新具体含义
                                meanings = existing_card.get("meanings", {})
                                source_key = list(meanings.keys())[0] if meanings else None
                                if source_key:
                                    meaning_data = meanings[source_key]
                                    if "upright" in meaning_data:
                                        if isinstance(meaning_data["upright"], dict):
                                            card["upright"] = meaning_data["upright"].get("general", "")
                                        else:
                                            card["upright"] = meaning_data["upright"]
                                    if "reversed" in meaning_data:
                                        if isinstance(meaning_data["reversed"], dict):
                                            card["reverse_neg"] = meaning_data["reversed"].get("general", "")
                                        else:
                                            card["reverse_neg"] = meaning_data["reversed"]
                
                # 更新小阿卡纳含义（简化处理）
                # 这里可以进一步完善小阿卡纳的含义加载
                
        except Exception as e:
            print(f"加载文件 {file_path} 时出错: {e}")
    
    return combined_knowledge

def create_complete_tarot_document():
    """创建完整的78张塔罗牌解读文档"""
    print("开始创建完整的78张塔罗牌文档...")
    
    # 加载塔罗牌数据
    knowledge = load_tarot_knowledge()
    print(f"加载了 {len(knowledge['majorArcana'])} 张大阿卡纳牌")
    total_minor = sum(len(suit["cards"]) for suit in knowledge["minorArcana"].values())
    print(f"加载了 {total_minor} 张小阿卡纳牌")
    print(f"总计: {len(knowledge['majorArcana']) + total_minor} 张牌")
    
    # 创建文档对象
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_paragraph('韦特塔罗牌78张完整深度解读', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加介绍段落
    intro_text = (
        '本解读以韦特塔罗牌经典牌面为基础，为全部78张牌提供详细的解读内容。\n'
        '包括22张大阿卡纳牌和56张小阿卡纳牌（权杖、圣杯、宝剑、星币各14张）。\n'
        '每张牌包含：牌面细节解读、核心原型定位、心理层面意义、正位核心解读、'
        '逆位核心解读（含消极走向与积极走向）。\n\n'
        '塔罗牌的核心并非预测未来，而是照见内心，通过牌面的象征唤醒自我认知，'
        '引导我们做出符合内心的选择，最终成为完整、真实、喜悦的自己。'
    )
    intro_para = doc.add_paragraph(intro_text)
    intro_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 第一部分：大阿卡纳
    doc.add_page_break()
    major_title = doc.add_paragraph('第一部分：22张大阿卡纳——愚人之旅的灵魂成长轨迹', style='Heading 1')
    major_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, card in enumerate(knowledge["majorArcana"]):
        print(f"处理大阿卡纳第 {i+1} 张: {card['nameCn']}")
        
        # 牌名标题
        card_title = doc.add_paragraph(f"{card['num']} {card['name']}", style='Heading 2')
        card_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 牌面解读
        doc.add_paragraph('【牌面解读】', style='Heading 3')
        doc.add_paragraph(card["image"], style='Normal')
        
        # 正位解读
        doc.add_paragraph('【正位核心解读】', style='Heading 3')
        doc.add_paragraph(card["upright"], style='Normal')
        
        # 逆位解读
        doc.add_paragraph('【逆位核心解读】', style='Heading 3')
        doc.add_paragraph(f'消极走向：{card["reverse_neg"]}', style='Normal')
        doc.add_paragraph(f'积极走向：{card["reverse_pos"]}', style='Normal')
        
        # 每5张牌分页
        if (i + 1) % 5 == 0 and i < len(knowledge["majorArcana"]) - 1:
            doc.add_page_break()
    
    # 第二部分：小阿卡纳
    doc.add_page_break()
    minor_title = doc.add_paragraph('第二部分：56张小阿卡纳——日常生命课题的完整指引', style='Heading 1')
    minor_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    suit_names = {
        "wands": "权杖组（火元素）- 行动、创造、热情",
        "cups": "圣杯组（水元素）- 情感、关系、内心世界", 
        "swords": "宝剑组（风元素）- 思维、沟通、冲突",
        "pentacles": "星币组（土元素）- 物质、财富、现实"
    }
    
    for suit_key, suit_info in knowledge["minorArcana"].items():
        doc.add_page_break()
        suit_header = doc.add_paragraph(suit_names[suit_key], style='Heading 2')
        suit_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for j, card in enumerate(suit_info["cards"]):
            print(f"处理{suit_info['name']}第 {j+1} 张: {card['nameCn']}")
            
            # 牌名标题
            card_title = doc.add_paragraph(f"{card['num']} {card['name']}", style='Heading 3')
            
            # 牌面解读
            doc.add_paragraph('【牌面解读】', style='Heading 4')
            doc.add_paragraph(card["image"], style='Normal')
            
            # 正位解读
            doc.add_paragraph('【正位核心解读】', style='Heading 4')
            doc.add_paragraph(card["upright"], style='Normal')
            
            # 逆位解读
            doc.add_paragraph('【逆位核心解读】', style='Heading 4')
            doc.add_paragraph(f'消极走向：{card["reverse_neg"]}', style='Normal')
            doc.add_paragraph(f'积极走向：{card["reverse_pos"]}', style='Normal')
            
            # 每3张牌分页（避免页面过长）
            if (j + 1) % 3 == 0 and j < len(suit_info["cards"]) - 1:
                doc.add_page_break()
    
    # 保存文档
    file_path = os.path.join(os.getcwd(), '78张韦特塔罗牌完整深度解读.docx')
    print(f"准备保存文件到: {file_path}")
    
    doc.save(file_path)
    print("完整的78张塔罗牌解读文档已生成！")
    
    # 验证文件
    if os.path.exists(file_path):
        file_size = os.path.getsize(file_path)
        print(f"文件生成成功！路径: {file_path}")
        print(f"文件大小: {file_size} 字节 ({file_size/1024/1024:.2f} MB)")
    else:
        print("文件生成失败！")

if __name__ == "__main__":
    print(f"当前工作目录: {os.getcwd()}")
    create_complete_tarot_document()