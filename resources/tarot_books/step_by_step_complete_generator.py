import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_complete_framework():
    """创建78张塔罗牌的完整框架结构"""
    print("🏗️ 开始构建78张塔罗牌完整框架...")
    
    # 创建文档
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 标题页
    title = doc.add_paragraph('韦特塔罗牌78张完整解读手册', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(
        '完整版详细解读指南\n'
        '涵盖22张大阿卡纳 + 56张小阿卡纳\n'
        '每张牌包含：牌面解读、原型定位、心理意义、正逆位详解',
        style='Subtitle'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 目录页
    doc.add_page_break()
    doc.add_paragraph('目录', style='Heading 1')
    
    # 大阿卡纳目录
    doc.add_paragraph('第一部分：大阿卡纳（22张）', style='Heading 2')
    major_names = [
        "0号 愚人", "1号 魔术师", "2号 女祭司", "3号 皇后", "4号 皇帝",
        "5号 教皇", "6号 恋人", "7号 战车", "8号 力量", "9号 隐士",
        "10号 命运之轮", "11号 正义", "12号 吊人", "13号 死神", 
        "14号 节制", "15号 恶魔", "16号 塔", "17号 星星", 
        "18号 月亮", "19号 太阳", "20号 审判", "21号 世界"
    ]
    
    for i, name in enumerate(major_names, 1):
        doc.add_paragraph(f'{i:2d}. {name}', style='Normal')
    
    # 小阿卡纳目录
    doc.add_paragraph('\n第二部分：小阿卡纳（56张）', style='Heading 2')
    suits = ["权杖组", "圣杯组", "宝剑组", "星币组"]
    
    start_num = 23
    for suit in suits:
        doc.add_paragraph(f'{suit}（14张）', style='Heading 3')
        for i in range(14):
            card_num = start_num + i
            if i < 10:
                card_name = f"{['一','二','三','四','五','六','七','八','九','十'][i]}号"
            else:
                court_names = ["侍卫", "骑士", "皇后", "国王"]
                card_name = court_names[i-10]
            doc.add_paragraph(f'{card_num:2d}. {suit[:-1]}{card_name}', style='Normal')
        start_num += 14
    
    # 开始添加详细内容
    doc.add_page_break()
    
    # 大阿卡纳详细内容模板
    print("📝 生成大阿卡纳详细内容...")
    doc.add_paragraph('第一部分：大阿卡纳详细解读', style='Heading 1')
    
    for i, name in enumerate(major_names):
        print(f"   处理第 {i+1:2d} 张: {name}")
        
        # 牌名标题
        doc.add_paragraph(f'{name}（The {name.split()[1]}）', style='Heading 2')
        
        # 各个解读部分占位符
        sections = [
            '【牌面细节解读】',
            '【核心原型定位】', 
            '【心理层面意义】',
            '【正位核心解读】',
            '【逆位核心解读】'
        ]
        
        for section in sections:
            doc.add_paragraph(section, style='Heading 3')
            # 这里预留空间供后续填充具体内容
            doc.add_paragraph('[详细内容待填充...]', style='Normal')
        
        # 每2张牌分页
        if (i + 1) % 2 == 0:
            doc.add_page_break()
    
    # 小阿卡纳详细内容模板
    print("📝 生成小阿卡纳详细内容...")
    doc.add_page_break()
    doc.add_paragraph('第二部分：小阿卡纳详细解读', style='Heading 1')
    
    suit_elements = [
        ("权杖", "火", "行动、创造、热情"),
        ("圣杯", "水", "情感、关系、内心"),
        ("宝剑", "风", "思维、沟通、冲突"),
        ("星币", "土", "物质、财富、现实")
    ]
    
    card_types = [f"{num}号" for num in ['一','二','三','四','五','六','七','八','九','十']] + ['侍卫', '骑士', '皇后', '国王']
    
    for suit_idx, (suit_name, element, theme) in enumerate(suit_elements):
        doc.add_page_break()
        doc.add_paragraph(f'{suit_name}组（{element}元素）- {theme}', style='Heading 2')
        
        for card_idx, card_type in enumerate(card_types):
            full_card_num = 23 + suit_idx * 14 + card_idx
            card_title = f'{full_card_num}号 {suit_name}{card_type}'
            
            print(f"   处理第 {full_card_num:2d} 张: {card_title}")
            
            doc.add_paragraph(card_title, style='Heading 3')
            
            # 小阿卡纳解读部分
            minor_sections = [
                '【元素特质】',
                '【牌面解读】',
                '【正位含义】',
                '【逆位含义】'
            ]
            
            for section in minor_sections:
                doc.add_paragraph(section, style='Heading 4')
                doc.add_paragraph('[详细内容待填充...]', style='Normal')
            
            # 每页3张牌
            if (card_idx + 1) % 3 == 0:
                doc.add_page_break()
    
    # 保存框架文档
    framework_filename = '78张韦特塔罗牌完整框架.docx'
    framework_path = os.path.join(os.getcwd(), framework_filename)
    
    doc.save(framework_path)
    print(f"✅ 框架文档已生成: {framework_path}")
    
    # 验证文件
    if os.path.exists(framework_path):
        size = os.path.getsize(framework_path)
        print(f"📊 文件大小: {size:,} 字节 ({size/1024:.1f} KB)")
        print(f"📈 预估页数: 约 {size/1500:.0f} 页")
    
    return framework_path

def main():
    print("🎯 韦特塔罗牌完整内容生成器")
    print("=" * 50)
    print(f"📍 工作目录: {os.getcwd()}")
    print("📋 本次任务: 生成78张塔罗牌完整框架文档")
    print("=" * 50)
    
    try:
        framework_path = create_complete_framework()
        print("=" * 50)
        print("✅ 78张塔罗牌完整框架文档生成完成！")
        print(f"📁 文件位置: {framework_path}")
        print("💡 下一步: 可在此框架基础上填充具体解读内容")
        print("=" * 50)
    except Exception as e:
        print(f"❌ 生成过程中出现错误: {e}")

if __name__ == "__main__":
    main()